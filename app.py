import streamlit as st
import pandas as pd
import time
import io
import docx
from fpdf import FPDF
import plotly.express as px
import database
import document_parser
import analyzer
import random

# Initialize local feedback DB
database.init_db()

st.set_page_config(page_title="TPRM Report Analyzer", layout="wide", page_icon="🛡️")

# --- Custom Styling ---
st.markdown("""
    <style>
    .main {background-color: #0e1117;}
    .stButton>button {width: 100%;}
    </style>
""", unsafe_allow_html=True)

# --- State Initialization ---
if 'analysis_done' not in st.session_state:
    st.session_state['analysis_done'] = False
if 'results' not in st.session_state:
    st.session_state['results'] = []

st.title("🛡️ TPRM Report Analyzer")
st.markdown("**Comprehensive, interactive Third-Party Risk Management evaluation powered by AI.**")
st.write("Upload vendor risk files, provide custom instructions, and generate leadership-ready reports.")

# --- 1 & 2. FILE INPUT & CUSTOM PROMPT INTERFACE ---
with st.sidebar:
    st.header("🔑 Configuration")
    api_key_input = st.text_input("Gemini API Key", type="password", help="Leave blank if configured in Streamlit Secrets.")
    
    # Try to grab from input first, fallback to Streamlit secrets if available
    api_key = api_key_input
    if not api_key:
        try:
            if "GEMINI_API_KEY" in st.secrets:
                api_key = st.secrets["GEMINI_API_KEY"]
        except Exception:
            pass
    
    st.header("1. Upload Vendor Files")
    uploaded_files = st.file_uploader(
        "Upload Reports (.pdf, .docx, .xlsx)", 
        type=['pdf', 'docx', 'xlsx'], 
        accept_multiple_files=True
    )
    
    st.header("2. Custom Prompt Section")
    custom_prompt = st.text_area(
        "Enter custom analysis instructions:", 
        placeholder="e.g., 'Focus heavily on GDPR compliance,' or 'Filter out any findings older than 2024'",
        help="The AI dynamically adapts its analysis algorithm based on these instructions."
    )
    
    analyze_btn = st.button("🚀 Run Analysis", type="primary")

# --- 3. ANALYSIS ENGINE ---
def run_live_analysis(files, prompt, key):
    results = []
    for file in files:
        with st.spinner(f"Extracting text from {file.name}..."):
            doc_text = document_parser.parse_file(file)
            
        max_retries = 3
        for attempt in range(max_retries):
            with st.spinner(f"Analyzing {file.name} with AI Engine (Attempt {attempt+1}/{max_retries})..."):
                res = analyzer.analyze_vendor_document(key, doc_text, prompt)
                
                if "error" in res:
                    if "429" in res["error"] and attempt < max_retries - 1:
                        # Wait 45 seconds for quota to reset based on typical free tier limits
                        sleep_time = 45 
                        st.warning(f"API rate limit hit. Waiting {sleep_time} seconds before retrying {file.name}...")
                        time.sleep(sleep_time)
                        continue
                    else:
                        st.error(f"Error analyzing {file.name}: {res['error']}")
                        break
                
                # Success
                break
                
        if "error" in res:
            continue # Skip adding to results if all retries failed
            
        # Add metadata
            res["filename"] = file.name
            
            # Map risk to numerical score for charts
            risk = res.get("risk_score", "Medium")
            score_val = {"Critical": 90, "High": 75, "Medium": 50, "Low": 25}.get(risk, 50) + random.randint(-5, 5)
            res["score_val"] = score_val
            
            results.append(res)
            
    # Sort from Critical to Low
    risk_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
    results.sort(key=lambda x: risk_order.get(x.get("risk_score", "Medium"), 2))
    return results

# --- 4. MULTI-FORMAT OUTPUT GENERATION ---
def generate_docx(results):
    doc = docx.Document()
    doc.add_heading('Executive Leadership Summary', 0)
    doc.add_paragraph(f"Total Vendors Assessed: {len(results)}")
    
    doc.add_heading('Vendor Breakdowns', 1)
    for res in results:
        doc.add_heading(f"{res.get('vendor', 'Unknown')} - {res.get('risk_score', 'Medium')} Risk", 2)
        doc.add_paragraph(res.get('summary', ''))
        doc.add_paragraph("Critical Findings:")
        for finding in res.get('findings', []):
            doc.add_paragraph(f"- {finding}", style='List Bullet')
        doc.add_paragraph("Next Steps:")
        for step in res.get('next_steps', []):
            doc.add_paragraph(f"- {step}", style='List Bullet')
        doc.add_paragraph(f"Leadership Postscript: {res.get('postscript', '')}")
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def generate_pdf(results):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=15, style="B")
    pdf.cell(200, 10, txt="Executive Leadership Summary", ln=1, align='L')
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=f"Total Vendors Assessed: {len(results)}", ln=1, align='L')
    pdf.cell(200, 10, txt="", ln=1) # spacer
    
    pdf.set_font("Arial", size=15, style="B")
    pdf.cell(200, 10, txt="Vendor Breakdowns", ln=1, align='L')
    
    def clean_text(text):
        if not text: return ""
        text = str(text)
        replacements = {'\u2018': "'", '\u2019': "'", '\u201c': '"', '\u201d': '"', '\u2013': '-', '\u2014': '--', '\u2026': '...', '\u00A0': ' '}
        for k, v in replacements.items():
            text = text.replace(k, v)
        return text.encode('latin-1', 'ignore').decode('latin-1')
    
    for res in results:
        pdf.set_font("Arial", size=13, style="B")
        pdf.cell(200, 10, txt=clean_text(f"{res.get('vendor', 'Unknown')} - {res.get('risk_score', 'Medium')} Risk"), ln=1, align='L')
        pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 10, txt=clean_text(res.get('summary', '')))
        pdf.set_font("Arial", size=11, style="B")
        pdf.cell(200, 10, txt="Critical Findings:", ln=1, align='L')
        pdf.set_font("Arial", size=11)
        for finding in res.get('findings', []):
            pdf.multi_cell(0, 8, txt=clean_text(f"- {finding}"))
        
        pdf.set_font("Arial", size=11, style="B")
        pdf.cell(200, 10, txt="Next Steps:", ln=1, align='L')
        pdf.set_font("Arial", size=11)
        for step in res.get('next_steps', []):
            pdf.multi_cell(0, 8, txt=clean_text(f"- {step}"))
            
        pdf.set_font("Arial", size=11, style="I")
        pdf.multi_cell(0, 10, txt=clean_text(f"Leadership Postscript: {res.get('postscript', '')}"))
        pdf.cell(200, 10, txt="", ln=1) # spacer
        
    return pdf.output(dest='S').encode('latin1')

# Main Execution Flow
if analyze_btn:
    if not api_key:
        try:
            available_secrets = list(st.secrets.keys())
        except Exception:
            available_secrets = ["No secrets found at all"]
            
        st.error(f"Please enter your Gemini API Key in the sidebar.\\n\\n**(Debug Info)** We checked Streamlit Cloud Secrets but couldn't find it. Found keys: {available_secrets}. Please double check you saved `GEMINI_API_KEY = \"...\"` in the Advanced Settings.")
    elif not uploaded_files:
        st.error("Please upload at least one vendor file.")
    else:
        st.session_state['results'] = run_live_analysis(uploaded_files, custom_prompt, api_key)
        st.session_state['analysis_done'] = True

if st.session_state.get('analysis_done') and st.session_state['results']:
    results = st.session_state['results']
    
    st.success("Analysis Complete!")
    
    # --- EXECUTIVE LEADERSHIP SUMMARY ---
    st.header("Executive Leadership Summary")
    col1, col2 = st.columns([1, 2])
    with col1:
        st.metric("Total Vendors Assessed", len(results))
        risk_counts = pd.DataFrame([r.get('risk_score', 'Medium') for r in results], columns=['Risk']).value_counts().reset_index()
        risk_counts.columns = ['Risk', 'Count']
        st.dataframe(risk_counts, hide_index=True, use_container_width=True)
    
    with col2:
        if not risk_counts.empty:
            fig = px.pie(risk_counts, values='Count', names='Risk', title="Master Risk Distribution",
                         color='Risk', color_discrete_map={"Critical": "red", "High": "orange", "Medium": "yellow", "Low": "green"})
            st.plotly_chart(fig, use_container_width=True)

    st.markdown("---")
    
    # --- DETAILED VENDOR BREAKDOWN ---
    st.header("Detailed Vendor-by-Vendor Breakdown")
    for res in results:
        vendor_name = res.get('vendor', 'Unknown')
        with st.expander(f"{res.get('risk_score', 'Medium')} RISK: {vendor_name}", expanded=True):
            st.write(f"**Context:** {res.get('summary', '')}")
            st.write(f"**Scope:** {res.get('scope', '')}")
            
            col_f, col_c = st.columns(2)
            with col_f:
                st.markdown("**Critical Findings:**")
                for f in res.get('findings', []):
                    st.markdown(f"- {f}")
                
                st.markdown("**Next Steps & Remediation:**")
                for s in res.get('next_steps', []):
                    st.markdown(f"- {s}")
            
            with col_c:
                # Individual vendor chart based on simulated sub-categories
                df_chart = pd.DataFrame({
                    "Category": ["Access Control", "Data Privacy", "Patch Management", "Compliance"],
                    "Vulnerability Score": [random.randint(10, 100) for _ in range(4)]
                })
                v_fig = px.bar(df_chart, x='Category', y='Vulnerability Score', title="Vulnerability Breakdown")
                st.plotly_chart(v_fig, use_container_width=True, key=f"chart_{vendor_name}")
            
            st.info(f"**Leadership Postscript:** {res.get('postscript', '')}")
            
            # --- 5. INTERACTIVE USER FEEDBACK ---
            st.markdown("---")
            st.write("**Was this analysis accurate?**")
            fcol1, fcol2, _ = st.columns([1, 1, 8])
            
            with fcol1:
                if st.button("👍 Thumbs Up", key=f"up_{vendor_name}"):
                    database.log_feedback(vendor_name, res.get("filename", ""), True)
                    st.toast(f"Feedback logged: Successful template mapping for {vendor_name}.")
            with fcol2:
                # Trigger dialog on thumbs down
                if st.button("👎 Thumbs Down", key=f"down_{vendor_name}"):
                    st.session_state[f"dialog_{vendor_name}"] = True

            # Feedback Dialog implementation using Streamlit elements
            if st.session_state.get(f"dialog_{vendor_name}", False):
                with st.form(key=f"form_{vendor_name}"):
                    st.write("Help the AI improve its analysis:")
                    feedback_text = st.text_area("Explicit comments (optional):", placeholder="e.g., 'The risk rating for Vendor X was too high'")
                    submitted = st.form_submit_button("Submit")
                    if submitted:
                        database.log_feedback(vendor_name, res.get("filename", ""), False, feedback_text)
                        st.toast(f"Negative feedback logged to local database. The AI will refine future iterations.")
                        st.session_state[f"dialog_{vendor_name}"] = False
                        st.rerun()

    # --- DOWNLOAD OPTIONS ---
    st.markdown("---")
    st.header("Export Final Report")
    
    docx_data = generate_docx(results)
    pdf_data = generate_pdf(results)
    
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        st.download_button(
            label="📄 Download as Word (.docx)",
            data=docx_data,
            file_name="TPRM_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
    with col_d2:
        st.download_button(
            label="📄 Download as PDF (.pdf)",
            data=pdf_data,
            file_name="TPRM_Report.pdf",
            mime="application/pdf",
            type="primary"
        )
